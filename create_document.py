from google.cloud import storage
from docx import Document
import io
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document(transcription, filename):
    file_stream = io.BytesIO()
    print("Creating Word document...")
    document = Document()

    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = 11

    for level in range(0,3):
        heading_style = document.styles[f"Heading {level}"]
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_style.border_bottom = None

    document.add_heading("Reading Title", 0)
    document.add_heading("A Commentary", level=2)
    document.add_paragraph(transcription)

    document.save(file_stream)
    file_stream.seek(0)
    print("Word document created.")
    
    client = storage.Client()
    bucket = client.bucket(os.getenv("GCS_TRANSCRIPTION_BUCKET_NAME"))

    docx_filename = f"{filename}.docx"
    blob = bucket.blob(docx_filename)

    blob.upload_from_file(file_stream, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document") 
    print("Word document uploaded to Cloud Storage.")
    file_stream.close()

    return f"<a href='{blob.public_url}'>Download the transcription here.</a>"