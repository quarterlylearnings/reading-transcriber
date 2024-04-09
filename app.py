from flask import Flask, request, send_file, jsonify
from google.cloud import speech, storage, firestore
from docx import Document
import os

app = Flask(__name__)

storage_client = storage.Client()
bucket_name = os.getenv("GCS_AUDIO_BUCKET_NAME")
bucket = storage_client.bucket(bucket_name)

db = firestore.Client()

@app.route("/upload", methods=["POST"])
def upload_file():
    print(request.files)
    if "file" not in request.files:
        return "No file part", 400
    file = request.files["file"]

    if file.filename == "":
        return "No selected file", 400

    if file:
        filename = file.filename

        blob = bucket.blob(filename)
        blob.upload_from_string(file.read(), content_type=file.content_type)

        gcs_path = f"gs://{bucket_name}/{filename}"
        store_file_metadata(filename, gcs_path)

        transcription = transcribe_audio(gcs_path)
        return create_word_document(transcription)
    
def store_file_metadata(filename, gcs_path, transcription_status="pending"):
    # log db object
    print(db.collection("reading-metadata"))
    doc_ref = db.collection("reading-metadata").document(filename)
    doc_ref.set({
        "filename": filename,
        "gcs_path": gcs_path,
        "transcription_status": transcription_status
    })
    
#TODO - PLACE IN A SEPARATE FILE (GOING TO GROW DUE TO ADDING CONTEXT)
def transcribe_audio(uri=None):
    client = speech.SpeechClient()

    audio = speech.RecognitionAudio(uri=uri)

    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=44100,
        language_code="en-US",
        enable_automatic_punctuation=True
    )

    operation = client.long_running_recognize(config=config, audio=audio)
    print("Waiting for operation to complete...")

    response = operation.result(timeout=90)

    transcription = ""
    for result in response.results:
        transcription += result.alternatives[0].transcript + "\n"

    return transcription

#TODO - PLACE IN A SEPARATE FILE
def create_word_document(transcription):
    document = Document()
    document.add_heading("Reading Title", 0)
    document.add_heading("A Commentary", level=2)
    document.add_paragraph(transcription)

    document.save("transcription.docx")

    return send_file("transcription.docx", as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)
